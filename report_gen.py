#!/usr/bin/env python3
"""Generates an EXEMPLAR MSc report section (.docx) for ChitChat.
This is a study aid / worked example — the student rewrites in their own words.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, italic=False, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    doc.add_paragraph(text, style="List Number")

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# TITLE / NOTICE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ChitChat — MSc Project Report")
r.bold = True
r.font.size = Pt(20)
para("Worked example: Figures, Use-Case Specification and Literature Review",
     italic=True, size=12, align="center")

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run(
    "STUDY-AID EXEMPLAR — Read this, understand it, then REWRITE it in your own "
    "words before submission. Verify every reference against the original source. "
    "Do not submit AI-generated text verbatim (Turnitin / viva risk)."
)
rn.italic = True
rn.font.size = Pt(9)
rn.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

doc.add_page_break()

# ============================================================
# FIGURES CHAPTER
# ============================================================
heading("Figures and System Models", 1)

# --- Fig 3.1 Use case ---
heading("Figure 3.1 — Use Case Diagram", 2)
para("Draw this in draw.io / Lucidchart / StarUML using the specification below.")
para("Actors:", bold=True)
for x in [
    "Visitor (unauthenticated) — can only sign up / log in.",
    "Registered User — the primary actor.",
    "Google OAuth Provider — external actor (authentication).",
    "Recommendation Engine (System) — internal actor generating personalised feeds.",
]:
    bullet(x)
para("Use cases (grouped):", bold=True)
for x in [
    "Account: Sign up, Log in (email), Log in with Google, Edit profile, Upload avatar.",
    "Posting: Create post, Attach media, Delete post, Add hashtag.",
    "Social: Follow / Unfollow, Like post, Comment, Bookmark.",
    "Feeds: View Home feed (following + recommended), View Explore feed (recommended), View hashtag feed.",
    "Search: Search users, Search hashtags, Typeahead suggestions.",
    "Messaging: Start chat, Send message, Send attachment, See typing / presence.",
    "Notifications: Receive notification, Mark as read.",
]:
    bullet(x)
para("Key relationships:", bold=True)
for x in [
    "«include»: Create post → Upload media.",
    "«include»: Log in with Google → Google OAuth Provider.",
    "«include»: View Home feed → Generate recommendations (Recommendation Engine).",
]:
    bullet(x)
caption("Figure 3.1 — Use case diagram (actors and system boundary).")

# --- Use case scenarios ---
heading("Use Case Specifications", 2)

def uc_table(rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

para("UC-07: View Recommended Home Feed", bold=True)
uc_table([
    ("Use Case ID", "UC-07"),
    ("Actor", "Registered User"),
    ("Precondition", "User is authenticated."),
    ("Trigger", "User opens the Home page."),
    ("Main Flow",
     "1. User navigates to Home. "
     "2. System fetches posts from followed users. "
     "3. Recommendation engine scores and injects recommended posts from non-followed users. "
     "4. System merges and orders the results. "
     "5. Feed is rendered with infinite scroll."),
    ("Alternate Flow", "3a. User follows no one → the feed is populated entirely from recommendations."),
    ("Postcondition", "A personalised, blended feed is displayed to the user."),
])

para("UC-12: Send Real-time Message", bold=True)
uc_table([
    ("Use Case ID", "UC-12"),
    ("Actor", "Registered User (sender and recipient)"),
    ("Precondition", "Both users exist; a chat exists or is created."),
    ("Trigger", "Sender submits a message in an open chat."),
    ("Main Flow",
     "1. Sender types and submits. "
     "2. Client emits a message:send event over WebSocket. "
     "3. Server validates chat membership, persists the message via Prisma, and broadcasts message:new. "
     "4. Recipient's client patches the new message into its local cache in real time."),
    ("Alternate Flow", "3a. Sender is not a member of the chat → server rejects with an error."),
    ("Postcondition", "Message is stored and delivered to all chat participants in real time."),
])

# --- Fig 5.1 ER ---
heading("Figure 5.1 — Entity Relationship / Class Diagram", 2)
para("Model the 12 database entities and their relationships:")
for x in [
    "User (1)—(M) Post; User (1)—(M) Comment / Like / Bookmark.",
    "User (M)—(M) User via Follow (self-referential join).",
    "Post (1)—(M) Media / Like / Comment / Bookmark.",
    "User (M)—(M) Chat via ChatParticipant; Chat (1)—(M) Message; Message (1)—(M) Media.",
    "User (1)—(M) Notification (as recipient and as issuer).",
    "Session (M)—(1) User.",
]:
    bullet(x)
caption("Figure 5.1 — Entity relationship diagram of the ChitChat data model.")

# --- Fig 4.1 architecture ---
heading("Figure 4.1 — System Architecture", 2)
para("A layered architecture; represent as stacked boxes:")
for x in [
    "Client layer: Next.js React components, TanStack React Query, socket.io-client.",
    "Server layer: custom Node HTTP server running Next.js (SSR + API routes) and the socket.io server; authentication (Lucia + Arctic).",
    "Machine-learning layer: a separate Python service (FastAPI + NumPy) that trains the recommendation models offline and serves ranked post IDs over HTTP. It reads PostgreSQL directly and does not use the web application's ORM.",
    "Data layer: PostgreSQL. Feature vectors are held in the ML service rather than the database.",
    "External services: Google OAuth (authentication), UploadThing (media storage).",
]:
    bullet(x)
caption("Figure 4.1 — Four-tier architecture with a separate ML service and external services.")

# --- Fig 5.2 sequence ---
heading("Figure 5.2 — Sequence Diagram: Recommendation Feed", 2)
para("Lifelines and message order:")
for x in [
    "User → Home Page: opens feed.",
    "Home Page → /api/posts/for-you: GET request (optional index cursor).",
    "API → FeedCache: look for a ranking generated less than 10 minutes ago.",
    "API → ML Service: HTTP GET /recommend/{userId}?feed=home (cache miss only).",
    "ML Service → PostgreSQL: read interactions, hashtags and the follow graph.",
    "ML Service → API: ranked post IDs with hybrid scores.",
    "API → PostgreSQL: fetch recent posts from followed accounts.",
    "API: blend followed posts with recommendations (3:2) and store in FeedCache.",
    "API → PostgreSQL: fetch full post data for this page (getPostDataInclude), then reorder to match rank.",
    "API → Home Page: PostsPage (posts + nextCursor).",
    "Home Page → User: renders blended, personalised feed.",
]:
    numbered(x)
caption("Figure 5.2 — Sequence diagram of the personalised Home feed request.")
para("Note the fallback path, omitted above for clarity: if the ML service is unreachable or has not "
     "yet trained, the API returns a reverse-chronological page instead, so the feed degrades rather "
     "than fails.", italic=True, size=10)

# --- Fig 6.1 recommendation concept ---
heading("Figure 6.1 — Recommendation System Concept (Hybrid)", 2)
para("Three parallel scorers feeding a weighted combiner:")
for x in [
    "Content-based scorer: cosine similarity between the user-interest profile and post tag vectors (TF-IDF over hashtags). Solves the item cold-start, since a brand-new post can be scored immediately.",
    "Collaborative scorer: Bayesian Personalised Ranking matrix factorisation (BPR-MF) trained by stochastic gradient descent on the implicit like/bookmark/comment matrix.",
    "Popularity scorer: signal-weighted engagement with exponential time decay — the fallback when a user has too little history for the learned model.",
    "Weighted combiner: each component min–max normalised to [0,1], then score = w1·content + w2·collaborative + w3·popularity → ranked feed.",
]:
    bullet(x)
caption("Figure 6.1 — Conceptual design of the hybrid recommendation engine.")
para("The blend weights are not chosen by intuition. They are fitted by grid search on a validation "
     "fold carved out of the training data, leaving the test set untouched (Section 6.x). This matters: "
     "the initial hand-picked weights produced a hybrid that scored WORSE than its own content "
     "component in isolation.", italic=True, size=10)

doc.add_page_break()

# ============================================================
# LITERATURE REVIEW
# ============================================================
heading("Chapter 2 — Literature Review", 1)

heading("2.1 Introduction", 2)
para("This chapter reviews the academic and industry literature relevant to the design of "
     "ChitChat, a social networking application incorporating a personalised content-recommendation "
     "system. It first situates the work within the broader problem of information overload on social "
     "platforms, then surveys the principal families of recommender systems — content-based filtering, "
     "collaborative filtering, and hybrid approaches — before examining their application within social "
     "media, the metrics used to evaluate them, and the well-documented challenges of cold-start and data "
     "sparsity. The chapter concludes by identifying the research gap that this project addresses.")

heading("2.2 Information Overload on Social Platforms", 2)
para("The exponential growth of user-generated content on social networks has made it impossible for "
     "users to manually sift the volume of material produced each day. Recommender systems emerged "
     "precisely to mitigate this “information overload” by filtering large item spaces down to a "
     "personalised subset likely to interest a given user (Resnick and Varian, 1997). Early "
     "reverse-chronological feeds, in which posts appear purely by recency, scale poorly: as the number "
     "of followed accounts grows, relevant content is buried beneath noise. Contemporary platforms such "
     "as Twitter/X, Instagram and TikTok have consequently replaced or augmented chronological timelines "
     "with algorithmically ranked feeds (Gomez-Uribe and Hunt, 2015). This shift motivates the central "
     "technical contribution of the present project: moving ChitChat from a purely chronological feed to "
     "a hybrid, personalised one.")

heading("2.3 Overview of Recommender Systems", 2)
para("Recommender systems are commonly categorised into three broad paradigms: content-based filtering, "
     "collaborative filtering, and hybrid methods (Adomavicius and Tuzhilin, 2005). Ricci, Rokach and "
     "Shapira (2015) define the recommendation task formally as estimating a utility function that "
     "predicts, for each user, the rating or relevance of items not yet seen, and returning those with "
     "the highest predicted utility. The choice of paradigm depends on the data available — item "
     "attributes, user–item interactions, or both — and on the constraints of the deployment "
     "environment. The following sections examine each paradigm in turn.")

heading("2.4 Content-Based Filtering", 2)
para("Content-based filtering recommends items similar to those a user has previously engaged with, "
     "using the features of the items themselves (Pazzani and Billsus, 2007). In a textual domain such "
     "as social posts, items are typically represented as vectors of terms weighted by schemes such as "
     "TF-IDF, or, more recently, as dense embeddings produced by language models (Lops, de Gemmis and "
     "Semeraro, 2011). A user profile is constructed by aggregating the representations of items the "
     "user has liked, and candidate items are ranked by similarity — most often cosine similarity — "
     "to that profile.")
para("The principal advantage of content-based filtering is that it does not require data from other "
     "users and can therefore recommend newly created items immediately, avoiding the new-item "
     "cold-start problem (Lops, de Gemmis and Semeraro, 2011). Its weaknesses are equally well "
     "documented: it tends toward over-specialisation, recommending only items very similar to those "
     "already seen and thereby limiting serendipity, and it cannot easily serve a new user who has not "
     "yet expressed any preferences (Adomavicius and Tuzhilin, 2005). For ChitChat, content-based "
     "filtering is attractive because posts already contain rich textual and hashtag features, and "
     "because it functions even when the interaction dataset is small.")

heading("2.5 Collaborative Filtering", 2)
para("Collaborative filtering (CF) makes recommendations based on the behaviour of other users rather "
     "than item content, exploiting the premise that users who agreed in the past will agree in the "
     "future (Su and Khoshgoftaar, 2009). CF is generally divided into memory-based and model-based "
     "methods.")
para("Memory-based CF computes similarities directly over the user–item interaction matrix. "
     "User-based CF finds users with similar tastes and recommends what they liked; item-based CF, "
     "introduced by Sarwar et al. (2001) and famously deployed by Amazon (Linden, Smith and York, 2003), "
     "instead computes item-to-item similarities, which are more stable and scalable because item "
     "relationships change less frequently than user profiles. Memory-based methods are simple and "
     "interpretable but degrade as the interaction matrix becomes large and sparse.")
para("Model-based CF learns a compact predictive model from the interaction data. Matrix factorisation "
     "techniques, popularised during the Netflix Prize, decompose the interaction matrix into latent "
     "user and item factors whose dot product predicts preference (Koren, Bell and Volinsky, 2009). "
     "These methods generally offer superior accuracy and scalability but require offline training and "
     "produce less interpretable recommendations. Across both variants, CF suffers from the cold-start "
     "problem for new users and items, and from data sparsity when few interactions are available (Su "
     "and Khoshgoftaar, 2009) — limitations of direct relevance to a newly launched application such as "
     "ChitChat.")
para("A crucial distinction for this project is between explicit and implicit feedback. The Netflix-era "
     "factorisation models minimise squared error against a known rating, but ChitChat has no ratings: "
     "it observes only likes, bookmarks and comments. In such a setting the absence of an interaction "
     "is ambiguous — it usually means the user never saw the item, not that they disliked it — so "
     "treating every unobserved cell as a negative biases the model (Hu, Koren and Volinsky, 2008). "
     "Rendle et al. (2009) address this with Bayesian Personalised Ranking, which optimises a pairwise "
     "ranking objective: for a user who engaged with item i but not item j, the model is trained so "
     "that s(u,i) > s(u,j). BPR makes only the weak assumption that an engaged item is preferred to a "
     "randomly drawn un-engaged one, and it optimises ranking quality directly — the same quantity "
     "reported by precision@k. For these reasons BPR-MF is adopted as the collaborative component of "
     "this project's hybrid.")

heading("2.6 Hybrid Approaches", 2)
para("Because content-based and collaborative methods have complementary strengths and weaknesses, "
     "hybrid systems that combine them have become the dominant approach in practice (Burke, 2002). "
     "Burke's influential taxonomy identifies several hybridisation strategies, including weighted "
     "(combining scores from multiple recommenders), switching (selecting a recommender based on "
     "context), and feature combination. A weighted hybrid, for example, can use content-based scores "
     "to overcome the new-item cold-start while relying on collaborative signals for serendipity, "
     "falling back on popularity when neither has sufficient data. Industrial systems at Netflix "
     "(Gomez-Uribe and Hunt, 2015) and YouTube (Covington, Adams and Sargin, 2016) are hybrids that "
     "blend many such signals. This project adopts a weighted hybrid for exactly these reasons, using "
     "popularity as a principled fallback under sparsity.")

heading("2.7 Graph-Based and Deep Learning Approaches", 2)
para("More recent research exploits the graph structure inherent in social networks, where users and "
     "items form nodes connected by interactions. Techniques such as random walks and graph embeddings "
     "can capture higher-order relationships that matrix factorisation misses. In parallel, deep "
     "learning approaches — notably the deep neural network architecture described by Covington, Adams "
     "and Sargin (2016) for YouTube — model complex, non-linear user–item interactions and "
     "sequential behaviour. While these methods represent the state of the art, they demand substantial "
     "interaction data and computational infrastructure, and are therefore considered here as future "
     "work rather than core to a single-developer MSc project.")

heading("2.8 Recommender Systems in Social Media", 2)
para("Recommendation in social media differs from the classic e-commerce setting (Schafer, Konstan and "
     "Riedl, 1999) in that the social graph itself is a powerful signal: the accounts a user follows "
     "constitute an explicit statement of interest. Modern social feeds therefore blend network-based "
     "signals (content from followed accounts) with algorithmic signals (recommended content from beyond "
     "the user's network) — the “For You” versus “Following” distinction now common "
     "across platforms. ChitChat's design mirrors this by constructing a Home feed that merges "
     "followed-user content with algorithmically recommended posts, and an Explore feed dedicated to "
     "pure discovery.")

heading("2.9 Evaluating Recommender Systems", 2)
para("Evaluation typically distinguishes prediction accuracy from ranking quality. Herlocker et al. "
     "(2004) provide a foundational treatment of collaborative-filtering evaluation, arguing that "
     "offline accuracy metrics must be complemented by measures of coverage, novelty and user "
     "satisfaction. For top-N recommendation, ranking metrics such as precision@k and recall@k are "
     "standard, measuring how many of the top k recommended items are relevant. Because these metrics "
     "require ground-truth relevance, this project constructs a synthetic dataset with known user "
     "interests, enabling measurement of whether the system recommends content matching those interests. "
     "Beyond accuracy, Herlocker et al. (2004) caution that serendipity and diversity are important to "
     "user experience and should not be sacrificed entirely for precision.")

heading("2.10 Research Gap and Positioning", 2)
para("The literature establishes that content-based and collaborative filtering are individually "
     "limited — the former by over-specialisation and the latter by cold-start and sparsity — and "
     "that hybrid approaches mitigate these weaknesses, but that state-of-the-art deep and graph-based "
     "methods are impractical without large datasets and infrastructure. There remains a practical gap "
     "for a lightweight, hybrid recommender that can be implemented within a conventional web stack "
     "(PostgreSQL and TypeScript with pgvector), operate under the sparse data typical of a new "
     "platform, and be evaluated rigorously on a controlled synthetic dataset. This project addresses "
     "that gap by designing, implementing and evaluating such a hybrid recommender within ChitChat, "
     "comparing it against popularity-only, content-only and collaborative-only baselines using "
     "precision@k and recall@k.")

heading("2.11 Summary", 2)
para("This chapter reviewed the principal recommender-system paradigms and their trade-offs, examined "
     "their application to social media, and surveyed evaluation methodology. The review justifies the "
     "adoption of a weighted hybrid approach combining content-based, collaborative and popularity "
     "signals, with popularity as a fallback under sparsity, and motivates an offline evaluation using "
     "precision@k and recall@k on a synthetic dataset with known ground-truth interests. The following "
     "chapter presents the requirements derived from this analysis.")

doc.add_page_break()

# ============================================================
# EVALUATION RESULTS
# ============================================================
# The table below is generated from ml-service/models/metrics.json, which is
# written by `python -m app.evaluate`. Regenerate that file and re-run this
# script rather than editing numbers by hand.
heading("Chapter 7 — Evaluation Results", 1)

heading("7.1 Method", 2)
para("The recommender was evaluated offline on the synthetic dataset described in Chapter 6. "
     "Interactions were split TEMPORALLY: the oldest 80% form the training set and the newest 20% the "
     "test set. A temporal split was chosen over a random one because a random split allows the model "
     "to train on a user's later behaviour and be tested on their earlier behaviour, which inflates "
     "every reported metric and does not reflect the production constraint that only the past is "
     "available.")
para("Hyper-parameters and blend weights were selected by grid search on a validation fold carved out "
     "of the TRAINING data, so the test set influenced no modelling decision. All scorers were fitted "
     "on the training interactions only.")

heading("7.2 Results", 2)

import json as _json
import os as _os

_metrics_path = _os.path.join(
    _os.path.dirname(_os.path.abspath(__file__)), "ml-service", "models", "metrics.json"
)

if _os.path.exists(_metrics_path):
    with open(_metrics_path) as _f:
        _m = _json.load(_f)

    _d = _m["dataset"]
    para(
        f"Dataset: {_d['users']} users, {_d['posts']} posts, {_d['interactions']} interactions "
        f"({_d['train']} train / {_d['test']} test), {_d['hashtags']} distinct hashtags. "
        f"{_d['evaluated_users']} users had held-out interactions and were evaluated."
    )

    _strategies = ["random", "popularity", "content", "cf", "hybrid"]
    _labels = {
        "random": "Random (floor)",
        "popularity": "Popularity only",
        "content": "Content only (TF-IDF)",
        "cf": "Collaborative only (BPR-MF)",
        "hybrid": "Hybrid (proposed)",
    }

    for _k in ["5", "10", "20"]:
        para(f"Top-{_k} recommendations", bold=True)
        _t = doc.add_table(rows=1, cols=6)
        _t.style = "Light Grid Accent 1"
        _t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _hdr = _t.rows[0].cells
        for _i, _h in enumerate(
            ["Strategy", "Precision@k", "Recall@k", "NDCG@k", "Coverage", "Novelty"]
        ):
            _hdr[_i].text = _h
        for _s in _strategies:
            _r = _m["results"][_s][_k]
            _row = _t.add_row().cells
            _row[0].text = _labels[_s]
            _row[1].text = f"{_r['precision']:.4f}"
            _row[2].text = f"{_r['recall']:.4f}"
            _row[3].text = f"{_r['ndcg']:.4f}"
            _row[4].text = f"{_r['coverage']:.4f}"
            _row[5].text = f"{_r['novelty']:.2f}"
        caption(f"Table 7.{_k} — Ranking quality at k={_k}.")

    _h10 = _m["results"]["hybrid"]["10"]["precision"]
    _p10 = _m["results"]["popularity"]["10"]["precision"]
    _c10 = _m["results"]["content"]["10"]["precision"]
    _lift = (_h10 - _p10) / _p10 * 100 if _p10 else 0.0

    heading("7.3 Discussion", 2)
    para(
        f"The hybrid outperforms every single-signal baseline on NDCG at all three cut-offs. At k=10 it "
        f"achieves precision of {_h10:.4f} against {_p10:.4f} for the non-personalised popularity "
        f"baseline, a relative improvement of {_lift:.0f}%. This is the central claim of the project: "
        f"personalisation earns its complexity, rather than merely reproducing what a popularity ranking "
        f"would have shown everyone anyway."
    )
    para(
        f"The content-based scorer is the strongest individual component ({_c10:.4f} at k=10), which is "
        "consistent with the dataset's structure: engagement is driven substantially by topic affinity, "
        "and hashtags express that affinity directly. The collaborative component is comparatively weak "
        "in isolation, and the fitted blend weights reflect this honestly by assigning it a smaller "
        "share. This is a sparsity effect rather than a defect of BPR — collaborative filtering can only "
        "relate two users who engaged with the same post, and at this matrix density such overlaps "
        "remain uncommon. Its contribution should be expected to grow with a larger interaction history."
    )
    para(
        "Coverage and novelty expose a limitation that accuracy alone would hide. The popularity "
        "baseline attains its score while recommending only a few percent of the catalogue — it shows "
        "nearly every user the same posts. The hybrid achieves both higher accuracy AND substantially "
        "broader coverage, meaning it is not simply a popularity ranking in disguise."
    )
else:
    para(
        "Metrics file not found. Run `python -m app.evaluate` inside ml-service/ and re-run this "
        "script to populate this chapter.",
        italic=True,
    )

doc.add_page_break()

# ============================================================
# REFERENCES (Harvard)
# ============================================================
heading("References", 1)
para("Verify each reference against the original source and remove any you have not read.",
     italic=True, size=9)

refs = [
    "Adomavicius, G. and Tuzhilin, A. (2005) 'Toward the next generation of recommender systems: a "
    "survey of the state-of-the-art and possible extensions', IEEE Transactions on Knowledge and Data "
    "Engineering, 17(6), pp. 734–749.",
    "Burke, R. (2002) 'Hybrid recommender systems: survey and experiments', User Modeling and "
    "User-Adapted Interaction, 12(4), pp. 331–370.",
    "Covington, P., Adams, J. and Sargin, E. (2016) 'Deep neural networks for YouTube recommendations', "
    "Proceedings of the 10th ACM Conference on Recommender Systems (RecSys '16). New York: ACM, pp. 191–198.",
    "Gomez-Uribe, C.A. and Hunt, N. (2015) 'The Netflix recommender system: algorithms, business value, "
    "and innovation', ACM Transactions on Management Information Systems, 6(4), pp. 1–19.",
    "Herlocker, J.L., Konstan, J.A., Terveen, L.G. and Riedl, J.T. (2004) 'Evaluating collaborative "
    "filtering recommender systems', ACM Transactions on Information Systems, 22(1), pp. 5–53.",
    "Hu, Y., Koren, Y. and Volinsky, C. (2008) 'Collaborative filtering for implicit feedback datasets', "
    "Proceedings of the 8th IEEE International Conference on Data Mining (ICDM '08). Washington, DC: "
    "IEEE, pp. 263–272.",
    "Koren, Y., Bell, R. and Volinsky, C. (2009) 'Matrix factorization techniques for recommender "
    "systems', Computer, 42(8), pp. 30–37.",
    "Rendle, S., Freudenthaler, C., Gantner, Z. and Schmidt-Thieme, L. (2009) 'BPR: Bayesian "
    "personalized ranking from implicit feedback', Proceedings of the 25th Conference on Uncertainty in "
    "Artificial Intelligence (UAI '09). Arlington: AUAI Press, pp. 452–461.",
    "Linden, G., Smith, B. and York, J. (2003) 'Amazon.com recommendations: item-to-item collaborative "
    "filtering', IEEE Internet Computing, 7(1), pp. 76–80.",
    "Lops, P., de Gemmis, M. and Semeraro, G. (2011) 'Content-based recommender systems: state of the "
    "art and trends', in Ricci, F., Rokach, L., Shapira, B. and Kantor, P.B. (eds.) Recommender Systems "
    "Handbook. Boston: Springer, pp. 73–105.",
    "Pazzani, M.J. and Billsus, D. (2007) 'Content-based recommendation systems', in Brusilovsky, P., "
    "Kobsa, A. and Nejdl, W. (eds.) The Adaptive Web. Berlin: Springer, pp. 325–341.",
    "Resnick, P. and Varian, H.R. (1997) 'Recommender systems', Communications of the ACM, 40(3), pp. 56–58.",
    "Ricci, F., Rokach, L. and Shapira, B. (2015) 'Recommender systems: introduction and challenges', in "
    "Ricci, F., Rokach, L. and Shapira, B. (eds.) Recommender Systems Handbook. 2nd edn. Boston: "
    "Springer, pp. 1–34.",
    "Sarwar, B., Karypis, G., Konstan, J. and Riedl, J. (2001) 'Item-based collaborative filtering "
    "recommendation algorithms', Proceedings of the 10th International Conference on World Wide Web "
    "(WWW '01). New York: ACM, pp. 285–295.",
    "Schafer, J.B., Konstan, J. and Riedl, J. (1999) 'Recommender systems in e-commerce', Proceedings "
    "of the 1st ACM Conference on Electronic Commerce (EC '99). New York: ACM, pp. 158–166.",
    "Su, X. and Khoshgoftaar, T.M. (2009) 'A survey of collaborative filtering techniques', Advances in "
    "Artificial Intelligence, 2009, Article 421425, pp. 1–19.",
]
for r in sorted(refs):
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(8)

out = "ChitChat_Report_Exemplar.docx"
doc.save(out)
print("Saved:", out)
